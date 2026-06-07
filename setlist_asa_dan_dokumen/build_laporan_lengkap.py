import docx
import json
import os
import sys
import re

# Ensure console supports UTF-8 on Windows
sys.stdout.reconfigure(encoding='utf-8')

# Determine absolute paths based on script directory to support execution from any CWD
script_dir = os.path.dirname(os.path.abspath(__file__))
RESULTS_FILE = os.path.join(script_dir, "plots/results.json")
TEMPLATE_FILE = os.path.join(script_dir, "output/CURRENT.docx")
OUTPUT_FILE = os.path.join(script_dir, "output/makalah_lengkap.docx")
PDF_FILE = os.path.join(script_dir, "output/makalah_lengkap.pdf")

PLOT_WAKTU = os.path.join(script_dir, "plots/perbandingan_waktu_plot.png")
PLOT_ENERGI = os.path.join(script_dir, "plots/kurva_energi_plot.png")
PLOT_KONVERGENSI = os.path.join(script_dir, "plots/konvergensi_aco_plot.png")

print("=" * 60)
print("COMPILER: Mengupdate Makalah Lengkap dari CURRENT.docx...")
print("=" * 60)

# Check requirements
if not os.path.exists(RESULTS_FILE):
    print(f"[ERROR] Hasil simulasi '{RESULTS_FILE}' tidak ditemukan.")
    print("        Silakan jalankan simulasi di src/main.py terlebih dahulu.")
    sys.exit(1)

if not os.path.exists(TEMPLATE_FILE):
    print(f"[ERROR] Template dokumen '{TEMPLATE_FILE}' tidak ditemukan.")
    sys.exit(1)

# Read results.json
with open(RESULTS_FILE, "r", encoding="utf-8") as f:
    results = json.load(f)

songs = results["songs"]
key_names = results["key_names"]
meta = results["meta"]

d_res = results["dijkstra"]
bb_res = results["branch_and_bound"]
aco_res = results["modified_aco"]

# Load CURRENT.docx
doc = docx.Document(TEMPLATE_FILE)

def set_cell_text_preserve_formatting(cell, text):
    """Sets text in cell while preserving formatting of the first run."""
    if not cell.paragraphs:
        cell.add_paragraph()
    p = cell.paragraphs[0]
    if not p.runs:
        run = p.add_run(text)
    else:
        run = p.runs[0]
        run.text = text
        for r in p.runs[1:]:
            r.text = ""
    for other_p in cell.paragraphs[1:]:
        other_p.text = ""

# Extract old times for text replacement before table update
old_d_time = doc.tables[4].rows[1].cells[2].text.strip()
old_bb_time = doc.tables[4].rows[2].cells[2].text.strip()
old_aco_time = doc.tables[4].rows[3].cells[2].text.strip()

old_ratio = int(round(float(old_bb_time) / float(old_aco_time))) if old_aco_time else 11
new_ratio = int(round(bb_res['time_ms'] / aco_res['time_ms']))
old_bb_sec = f"~{float(old_bb_time)/1000:.1f} detik" if old_bb_time else "~6.8 detik"
new_bb_sec = f"~{bb_res['time_ms']/1000:.1f} detik"

# 1. Update Table 4 (Comparison Table)
print("Updating Table 4 (Tabel Perbandingan)...")
table_4 = doc.tables[4]
set_cell_text_preserve_formatting(table_4.rows[1].cells[1], f"{d_res['cost']:.2f}")
set_cell_text_preserve_formatting(table_4.rows[1].cells[2], f"{d_res['time_ms']:.4f}")
set_cell_text_preserve_formatting(table_4.rows[1].cells[3], f"{d_res['drops']}x")
set_cell_text_preserve_formatting(table_4.rows[1].cells[4], f"{d_res['drops'] * meta['lambda']:.2f}")
set_cell_text_preserve_formatting(table_4.rows[1].cells[5], f"{d_res['cost'] - d_res['drops'] * meta['lambda']:.2f}")

set_cell_text_preserve_formatting(table_4.rows[2].cells[1], f"{bb_res['cost']:.2f}")
set_cell_text_preserve_formatting(table_4.rows[2].cells[2], f"{bb_res['time_ms']:.2f}")
set_cell_text_preserve_formatting(table_4.rows[2].cells[3], f"{bb_res['drops']}x")
set_cell_text_preserve_formatting(table_4.rows[2].cells[4], f"{bb_res['drops'] * meta['lambda']:.2f}")
set_cell_text_preserve_formatting(table_4.rows[2].cells[5], f"{bb_res['cost'] - bb_res['drops'] * meta['lambda']:.2f}")

set_cell_text_preserve_formatting(table_4.rows[3].cells[1], f"{aco_res['cost']:.2f}")
set_cell_text_preserve_formatting(table_4.rows[3].cells[2], f"{aco_res['time_ms']:.2f}")
set_cell_text_preserve_formatting(table_4.rows[3].cells[3], f"{aco_res['drops']}x")
set_cell_text_preserve_formatting(table_4.rows[3].cells[4], f"{aco_res['drops'] * meta['lambda']:.2f}")
set_cell_text_preserve_formatting(table_4.rows[3].cells[5], f"{aco_res['cost'] - aco_res['drops'] * meta['lambda']:.2f}")

# 2. Update Table 5 (Dijkstra Setlist)
print("Updating Table 5 (Dijkstra Setlist)...")
table_5 = doc.tables[5]
for pos, idx in enumerate(d_res["path"]):
    row_idx = pos + 1
    s = songs[idx]
    drop = pos > 0 and s["energy"] < songs[d_res["path"][pos - 1]]["energy"]
    status = "Pembuka" if pos == 0 else ("Energy Drop (Penalti)" if drop else "Ascending")
    col_values = [
        str(pos + 1),
        s["title"],
        f"{s['energy']}/10",
        str(s["bpm"]),
        f"{key_names[s['key']]} ({s['key']})",
        status
    ]
    for col_idx, val in enumerate(col_values):
        cell = table_5.rows[row_idx].cells[col_idx]
        set_cell_text_preserve_formatting(cell, val)
        p = cell.paragraphs[0]
        if p.runs:
            run = p.runs[0]
            if drop:
                run.font.color.rgb = docx.shared.RGBColor(239, 68, 68)
                run.font.bold = True
            else:
                run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
                run.font.bold = False

# 3. Update Table 6 (Branch and Bound Setlist)
print("Updating Table 6 (Branch and Bound Setlist)...")
table_6 = doc.tables[6]
for pos, idx in enumerate(bb_res["path"]):
    row_idx = pos + 1
    s = songs[idx]
    drop = pos > 0 and s["energy"] < songs[bb_res["path"][pos - 1]]["energy"]
    status = "Pembuka" if pos == 0 else ("Energy Drop (Penalti)" if drop else "Ascending")
    col_values = [
        str(pos + 1),
        s["title"],
        f"{s['energy']}/10",
        str(s["bpm"]),
        f"{key_names[s['key']]} ({s['key']})",
        status
    ]
    for col_idx, val in enumerate(col_values):
        cell = table_6.rows[row_idx].cells[col_idx]
        set_cell_text_preserve_formatting(cell, val)
        p = cell.paragraphs[0]
        if p.runs:
            run = p.runs[0]
            if drop:
                run.font.color.rgb = docx.shared.RGBColor(239, 68, 68)
                run.font.bold = True
            else:
                run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
                run.font.bold = False

# 4. Update Table 7 (Cost Distribution Table)
print("Updating Table 7 (Tabel Distribusi Cost)...")
table_7 = doc.tables[7]
# Dijkstra
set_cell_text_preserve_formatting(table_7.rows[1].cells[1], str(d_res['drops']))
set_cell_text_preserve_formatting(table_7.rows[1].cells[2], f"{d_res['drops'] * meta['lambda']:.2f}")
set_cell_text_preserve_formatting(table_7.rows[1].cells[3], f"{d_res['cost'] - d_res['drops'] * meta['lambda']:.2f}")
set_cell_text_preserve_formatting(table_7.rows[1].cells[4], f"{d_res['cost']:.2f}")

# B&B
set_cell_text_preserve_formatting(table_7.rows[2].cells[1], str(bb_res['drops']))
set_cell_text_preserve_formatting(table_7.rows[2].cells[2], f"{bb_res['drops'] * meta['lambda']:.2f}")
set_cell_text_preserve_formatting(table_7.rows[2].cells[3], f"{bb_res['cost'] - bb_res['drops'] * meta['lambda']:.2f}")
set_cell_text_preserve_formatting(table_7.rows[2].cells[4], f"{bb_res['cost']:.2f}")

# ACO
set_cell_text_preserve_formatting(table_7.rows[3].cells[1], str(aco_res['drops']))
set_cell_text_preserve_formatting(table_7.rows[3].cells[2], f"{aco_res['drops'] * meta['lambda']:.2f}")
set_cell_text_preserve_formatting(table_7.rows[3].cells[3], f"{aco_res['cost'] - aco_res['drops'] * meta['lambda']:.2f}")
set_cell_text_preserve_formatting(table_7.rows[3].cells[4], f"{aco_res['cost']:.2f}")

# 5. Update Table 8 (ACO Convergence Log)
print("Updating Table 8 (ACO Convergence Log)...")
table_8 = doc.tables[8]
log = aco_res["convergence_log"]
set_cell_text_preserve_formatting(table_8.rows[1].cells[1], f"{log[0]:.2f}")
set_cell_text_preserve_formatting(table_8.rows[2].cells[1], f"{log[1]:.2f}")
set_cell_text_preserve_formatting(table_8.rows[3].cells[1], f"{log[2]:.2f}")
set_cell_text_preserve_formatting(table_8.rows[4].cells[1], f"{log[3]:.2f}")
set_cell_text_preserve_formatting(table_8.rows[5].cells[1], f"{log[-1]:.2f}")

# 6. Dynamic search-and-replace in paragraph runs for times/ratios
print("Mengupdate tulisan hasil eksekusi dalam paragraf...")
p_indices = [145, 157, 191]
for idx in p_indices:
    p = doc.paragraphs[idx]
    for run in p.runs:
        if old_aco_time in run.text:
            run.text = run.text.replace(old_aco_time, f"{aco_res['time_ms']:.2f}")
        if old_bb_time in run.text:
            run.text = run.text.replace(old_bb_time, f"{bb_res['time_ms']:.2f}")
        if old_d_time in run.text:
            run.text = run.text.replace(old_d_time, f"{d_res['time_ms']:.4f}")
        # Replace ratio
        run.text = re.sub(r'(\d+)\s+kali', f"{new_ratio} kali", run.text)
        # Replace seconds
        if old_bb_sec in run.text:
            run.text = run.text.replace(old_bb_sec, new_bb_sec)

# 7. Replace plot images in paragraphs 151, 170, 183
def replace_image_in_paragraph(p_idx, img_path):
    if not os.path.exists(img_path):
        print(f"[WARNING] File plot '{img_path}' tidak ditemukan. Melewati update gambar.")
        return
    p = doc.paragraphs[p_idx]
    p.text = ""  # Clear old runs & drawing OLE/image
    run = p.add_run()
    run.add_picture(img_path, width=docx.shared.Inches(5.5))
    p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

print("Mengupdate gambar grafik visualisasi...")
replace_image_in_paragraph(151, PLOT_WAKTU)
replace_image_in_paragraph(170, PLOT_ENERGI)
replace_image_in_paragraph(183, PLOT_KONVERGENSI)

# Save output docx
try:
    doc.save(OUTPUT_FILE)
    print(f"[OK] Berkas DOCX berhasil disimpan ke: {OUTPUT_FILE}")
except PermissionError:
    alt_output = OUTPUT_FILE.replace(".docx", "_baru.docx")
    doc.save(alt_output)
    print(f"[WARNING] Gagal menulis ke '{OUTPUT_FILE}' karena file sedang dibuka.")
    print(f"          Berkas alternatif disimpan ke: {alt_output}")
    OUTPUT_FILE = alt_output

# 8. Convert to PDF
def convert_to_pdf(docx_path, pdf_path):
    print("Mencoba membuat PDF via Word COM Automation...")
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc_obj = word.Documents.Open(os.path.abspath(docx_path))
        doc_obj.SaveAs(os.path.abspath(pdf_path), FileFormat=17) # 17 = wdFormatPDF
        doc_obj.Close()
        word.Quit()
        print(f"[OK] Berkas PDF berhasil disimpan ke: {pdf_path}")
        return True
    except Exception as e:
        print(f"[INFO] COM Automation gagal: {e}")
        # Fallback to LibreOffice if soffice is available in PATH
        print("Mencoba membuat PDF via LibreOffice...")
        try:
            import subprocess
            subprocess.run(["soffice", "--headless", "--convert-to", "pdf", "--outdir", os.path.dirname(pdf_path), docx_path], check=True)
            print(f"[OK] Berkas PDF berhasil disimpan ke: {pdf_path}")
            return True
        except Exception as le:
            print(f"[INFO] LibreOffice gagal: {le}")
            print("\n[PERINGATAN MANUAL]")
            print(f"  Berkas PDF '{pdf_path}' tidak dapat dibuat secara otomatis.")
            print(f"  Silakan buka berkas Word '{docx_path}'")
            print("  dan simpan secara manual sebagai PDF (Save As -> PDF) di komputer Anda.")
            return False

convert_to_pdf(OUTPUT_FILE, PDF_FILE)
print("=" * 60)
print("COMPILER SELESAI!")
print("=" * 60)
